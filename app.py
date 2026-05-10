import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sqlite3
import os
import zipfile
import io
from datetime import date


# --- CONFIGURATION DE LA PAGE ---
st.set_page_config(
    page_title="CATI - Générateur d'Attestations",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- COULEURS ET STYLE CSS ---
CATI_BLUE = "#1E3A8A"
CATI_GREEN = "#047857"

st.markdown(f"""
    <style>
    .reportview-container {{ background-color: #F9FAFB; }}
    h1, h2, h3 {{ color: {CATI_BLUE}; }}
    .stButton>button {{ background-color: {CATI_GREEN}; color: white; border-radius: 8px; }}
    .stButton>button:hover {{ background-color: #065F46; color: white; }}
    </style>
    """, unsafe_allow_html=True)

# --- INITIALISATION BDD ---
DB_DIR = "data"
DB_PATH = os.path.join(DB_DIR, "attestations.db")

def init_db():
    if not os.path.exists(DB_DIR):
        os.makedirs(DB_DIR)
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS history
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  reference TEXT, participant_name TEXT, formation_title TEXT, date_gen TEXT)''')
    conn.commit()
    conn.close()

def save_to_db(reference, participant_name, formation_title, date_gen):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("INSERT INTO history (reference, participant_name, formation_title, date_gen) VALUES (?, ?, ?, ?)",
              (reference, participant_name, formation_title, date_gen))
    conn.commit()
    conn.close()

def get_history():
    conn = sqlite3.connect(DB_PATH)
    df = pd.read_sql_query("SELECT * FROM history ORDER BY id DESC", conn)
    conn.close()
    return df

init_db()

# --- LOGIQUE DE GÉNÉRATION WORD ---

def remove_table_borders(table):
    """Supprime toutes les bordures d'un tableau Word"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        borders.append(border)
    tblPr.append(borders)

def add_header_table(doc, logo_left_path, logo_right_path):
    """Ajoute l'en-tête avec les logos et le texte officiel hierarchisé"""
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)
    
    for cell in table.columns[0].cells:
        cell.width = Cm(3.5)
    for cell in table.columns[1].cells:
        cell.width = Cm(12)
    for cell in table.columns[2].cells:
        cell.width = Cm(3.5)

    cell_left = table.cell(0, 0)
    cell_left.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(logo_left_path):
        p_left = cell_left.paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_left.add_run()
        run.add_picture(logo_left_path, width=Inches(0.75))

    cell_center = table.cell(0, 1)
    cell_center.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p1 = cell_center.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.space_after = Pt(0)
    p1.space_before = Pt(0)
    run1 = p1.add_run("République Algérienne Démocratique et Populaire")
    run1.font.size = Pt(9)
    run1.font.name = 'Times New Roman'
    
    p2 = cell_center.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.space_after = Pt(0)
    p2.space_before = Pt(0)
    run2 = p2.add_run("Ministère De L’Enseignement Et De La Recherche Scientifique")
    run2.font.size = Pt(9)
    run2.font.name = 'Times New Roman'
    
    p3 = cell_center.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.space_after = Pt(2)
    p3.space_before = Pt(2)
    run3 = p3.add_run("Centre Universitaire De Maghnia")
    run3.font.size = Pt(11)
    run3.bold = True
    run3.font.name = 'Times New Roman'
    
    p4 = cell_center.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.space_after = Pt(0)
    p4.space_before = Pt(4)
    run4 = p4.add_run("Centre D’Appui De La Technologie Et De L’Innovation")
    run4.font.size = Pt(12)
    run4.bold = True
    run4.font.name = 'Times New Roman'
    
    p5 = cell_center.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.space_after = Pt(0)
    p5.space_before = Pt(0)
    run5 = p5.add_run("-CATI-")
    run5.font.size = Pt(12)
    run5.bold = True
    run5.font.name = 'Times New Roman'

    cell_right = table.cell(0, 2)
    cell_right.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(logo_right_path):
        p_right = cell_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_right.add_run()
        run.add_picture(logo_right_path, width=Inches(0.75))


def generate_attestation_docx(data, participant):
    doc = Document()

    sections = doc.sections
    for section in sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.header_distance = Cm(0.5)
        section.footer_distance = Cm(0.5)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    add_header_table(doc, "assets/logo_ministry.png", "assets/logo_university.png")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(data["type_attestation"].upper())
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    p_ref = doc.add_paragraph()
    p_ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ref = p_ref.add_run(f"Référence : {data['reference']}")
    run_ref.font.size = Pt(10)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)

    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_after = Pt(6)
    p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_intro.add_run("Nous soussignés, le Centre d’Appui à la Technologie et à l’Innovation (CATI) de Maghnia, certifions que :")

    # Sécurisation des données du participant pour éviter l'impression de "nan"
    genre = str(participant["Genre"]) if pd.notna(participant["Genre"]) else ""
    nom = str(participant["Nom et Prénom"]) if pd.notna(participant["Nom et Prénom"]) else "NOM PRENOM"
    qualite = str(participant["Qualité"]) if pd.notna(participant["Qualité"]) else ""

    p_part = doc.add_paragraph()
    p_part.paragraph_format.space_after = Pt(6)
    p_part.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_part = p_part.add_run(f"{genre} {nom}")
    run_part.bold = True
    run_part.font.size = Pt(14)

    p_qual = doc.add_paragraph()
    p_qual.paragraph_format.space_after = Pt(6)
    p_qual.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_qual = p_qual.add_run(f"Qualité : {qualite}")
    run_qual.italic = True

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)

    p_suite = doc.add_paragraph()
    p_suite.paragraph_format.space_after = Pt(6)
    p_suite.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_suite.add_run("A suivi avec assiduité la formation spécialisée intitulée :")

    doc.add_paragraph() 

    p_titre_form = doc.add_paragraph()
    p_titre_form.paragraph_format.space_after = Pt(6)
    p_titre_form.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tf = p_titre_form.add_run(f"« {data['titre_formation']} »")
    run_tf.bold = True
    run_tf.font.size = Pt(13)

    doc.add_paragraph() 

    p_details_title = doc.add_paragraph()
    p_details_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_dt = p_details_title.add_run("Détails de la session :")
    run_dt.bold = True

    details_lines = [
        f"Date : {data['date_formation']}",
        f"Durée totale : {data['duree']}",
        f"Lieu : {data['lieu']}"
    ]

    for line in details_lines:
        p_detail = doc.add_paragraph(line)
        p_detail.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_detail.paragraph_format.space_after = Pt(2)
        p_detail.paragraph_format.space_before = Pt(0)

    if data.get('formateur'):
        p_form = doc.add_paragraph()
        p_form.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_form.paragraph_format.space_after = Pt(6)
        run_form_label = p_form.add_run("Formateur(s) : ")
        run_form_label.bold = True
        run_form_value = p_form.add_run(data['formateur'])

    doc.add_paragraph() 

    p_obj_title = doc.add_paragraph()
    p_obj_title.paragraph_format.space_after = Pt(6)
    p_obj_title.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_ot = p_obj_title.add_run("Objectifs et compétences acquises : ")
    run_ot.bold = True

    participe = "la participante" if genre in ["Mme", "Mlle"] else "le participant"
    p_obj_title.add_run(f"Au cours de cette formation, {participe} a acquis une expertise pratique sur les thématiques suivantes :")

    objectifs = [obj.strip() for obj in data["objectifs"].split('\n') if obj.strip()]
    for obj in objectifs:
        if obj[0].isdigit():
            obj = '.'.join(obj.split('.')[1:]).strip()
        p_obj = doc.add_paragraph(obj, style='List Bullet')
        p_obj.paragraph_format.space_after = Pt(3)
        p_obj.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if data.get("evaluation"):
        p_eval = doc.add_paragraph()
        p_eval.paragraph_format.space_after = Pt(6)
        p_eval.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_eval.add_run("L'évaluation finale a permis de valider les acquis théoriques et les exercices pratiques d'application sur plateforme.")

    p_close = doc.add_paragraph()
    p_close.paragraph_format.space_after = Pt(6)
    p_close.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_close.add_run("En foi de quoi, la présente attestation lui est délivrée pour servir et valoir ce que de droit.")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)

    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_sig = p_sig.add_run(f"Le {data['poste_responsable']}")
    run_sig.bold = True

    p_sig_nom = doc.add_paragraph()
    p_sig_nom.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sig_nom.add_run(data["nom_responsable"])

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# --- INTERFACE UTILISATEUR ---

with st.sidebar:
    st.image("assets/logo_cati.png" if os.path.exists("assets/logo_cati.png") else "https://via.placeholder.com/150x150?text=CATI", width=150)
    st.title("⚙️ Paramètres")
    st.markdown("### Modèle d'attestation")
    model_choice = st.selectbox("Choisir le modèle", ["Modèle CATI", "Modèle Université (Brouillon)"])

st.title("🎓 Générateur d'Attestations de Formation")
st.markdown("---")

tab1, tab2, tab3 = st.tabs(["📝 Saisie des informations", "👁️ Aperçu & Génération", "📁 Historique"])

with tab1:
    with st.form("attestation_form"):
        st.header("1️⃣ Informations générales")
        col1, col2, col3 = st.columns(3)
        with col1:
            reference = st.text_input("Référence *", value="CATI/2026-01")
        with col2:
            type_attestation = st.selectbox("Type d'attestation *", ["Attestation de formation", "Attestation de participation", "Certificat de présence"])
        with col3:
            date_generation = st.date_input("Date de génération *", value=date.today())

        st.header("👥 2️⃣ Gestion des participants")
        
        st.markdown("##### 📥 Importer une liste de participants")
        uploaded_file = st.file_uploader("Chargez un fichier Excel (.xlsx) ou CSV contenant les colonnes : `Nom et Prénom`, `Genre`, `Qualité`, `Spécialité`, `Institution`", type=['csv', 'xlsx'], key="uploader_participants")
        
        expected_cols = ["Nom et Prénom", "Genre", "Qualité", "Spécialité", "Institution"]
        if "participants_df" not in st.session_state:
            st.session_state.participants_df = pd.DataFrame(columns=expected_cols)

        if uploaded_file is not None:
            try:
                if uploaded_file.name.endswith('.csv'):
                    df_uploaded = pd.read_csv(uploaded_file)
                else:
                    df_uploaded = pd.read_excel(uploaded_file)
                
                for col in expected_cols:
                    if col not in df_uploaded.columns:
                        df_uploaded[col] = ""
                
                df_uploaded = df_uploaded[expected_cols]
                st.session_state.participants_df = df_uploaded
                st.success(f"✅ {len(df_uploaded)} participants importés avec succès !")
            except Exception as e:
                st.error(f"❌ Erreur lors de la lecture du fichier : {e}.")

        st.markdown("##### ✏️ Modifier / Ajouter manuellement")
        st.info("Vous pouvez modifier les cellules directement, ou ajouter/supprimer des lignes avec les boutons à droite du tableau.")

        participants_edited = st.data_editor(
            st.session_state.participants_df, 
            num_rows="dynamic",
            use_container_width=True
        )

        st.header("📚 3️⃣ Détails de la formation")
        titre_formation = st.text_input("Titre de la formation *", value="Maîtrise de la Recherche Brevet : Stratégies par Mots-Clés et Classification Internationale des Brevets (IPC) sur PATENTSCOPE (WIPO)")

        col4, col5, col6 = st.columns(3)
        with col4:
            date_formation = st.text_input("Date de la formation *", value="08 mai 2026")
        with col5:
            duree = st.text_input("Durée *", value="03 heures")
        with col6:
            lieu = st.text_input("Lieu *", value="Salle 24 ST, Centre Universitaire de Maghnia")

        st.markdown("##### 👨‍🏫 Formateur(s)")
        formateur = st.text_input("Nom du ou des formateur(s)", value="", help="Laissez vide si non applicable. Séparez par des virgules s'il y en a plusieurs.", placeholder="Ex: Dr. Ahmed Benali, Pr. Dupont")

        st.header("✅ 4️⃣ Objectifs et compétences acquises")
        objectifs = st.text_area("Objectifs (un par ligne) *", value="1. Exploitation de la base PATENTSCOPE : Navigation avancée\n2. Maîtrise de la Classification Internationale des Brevets (IPC)\n3. Analyse de l'état de la technique", height=150)
        evaluation = st.checkbox("Évaluation finale validée", value=True)

        st.header("📄 5️⃣ Signature et validation")
        col7, col8 = st.columns(2)
        with col7:
            nom_responsable = st.text_input("Nom du responsable *", value="Dr. FODIL Mohammed El Amine")
        with col8:
            poste_responsable = st.text_input("Poste du responsable *", value="Responsable de la Formation")
        
        submitted = st.form_submit_button("✅ Valider et Prévisualiser")
        if submitted:
            if not reference or not titre_formation or not objectifs:
                st.error("Veuillez remplir tous les champs obligatoires (*)")
            elif participants_edited.empty:
                st.error("Veuillez ajouter au moins un participant.")
            else:
                st.session_state.form_validated = True
                st.session_state.form_data = {
                    "reference": reference,
                    "type_attestation": type_attestation,
                    "date_generation": str(date_generation),
                    "titre_formation": titre_formation,
                    "date_formation": date_formation,
                    "duree": duree,
                    "lieu": lieu,
                    "formateur": formateur,
                    "objectifs": objectifs,
                    "evaluation": evaluation,
                    "nom_responsable": nom_responsable,
                    "poste_responsable": poste_responsable
                }
                st.session_state.participants_df = participants_edited
                st.success("Formulaire validé ! Allez dans l'onglet 'Aperçu & Génération'.")

with tab2:
    if 'form_validated' in st.session_state and st.session_state.form_validated:
        data = st.session_state.form_data
        df_participants = st.session_state.participants_df
        
        st.header("Aperçu de l'attestation")
        
        # --- CORRECTION DU BUG DU SELECTBOX ---
        def format_participant_name(i):
            val = df_participants.iloc[i]["Nom et Prénom"]
            if pd.isna(val) or str(val).strip() == "":
                return f"Participant {i+1} (Nom manquant)"
            return str(val)

        selected_participant_idx = st.selectbox(
            "Choisir un participant pour l'aperçu", 
            range(len(df_participants)), 
            format_func=format_participant_name
        )
        
        participant = df_participants.iloc[selected_participant_idx]
        
        formateur_html = f"<b>Formateur(s) :</b> {data['formateur']}<br>" if data.get('formateur') else ""
        
        # Sécurisation pour l'aperçu HTML aussi
        p_genre = str(participant['Genre']) if pd.notna(participant['Genre']) else ""
        p_nom = str(participant['Nom et Prénom']) if pd.notna(participant['Nom et Prénom']) else "NOM PRENOM"
        p_qualite = str(participant['Qualité']) if pd.notna(participant['Qualité']) else ""

        st.markdown(f"""
        <div style="border: 1px solid #ccc; padding: 20px; border-radius: 5px; background-color: white; font-family: Times New Roman;">
            <p style="text-align: center; font-size: 9px;">République Algérienne Démocratique et Populaire<br>Ministère De L’Enseignement Et De La Recherche Scientifique</p>
            <p style="text-align: center; font-weight: bold; font-size: 11px;">Centre Universitaire De Maghnia</p>
            <p style="text-align: center; font-weight: bold; font-size: 12px; color: {CATI_BLUE};">Centre D’Appui De La Technologie Et De L’Innovation<br>-CATI-</p>
            <hr>
            <p style="text-align: center; font-weight: bold; color: {CATI_BLUE}; font-size: 18px;">{data['type_attestation'].upper()}</p>
            <p style="text-align: center; font-size: 10px;">Référence : {data['reference']}</p>
            <hr>
            <p>Nous soussignés, le Centre d’Appui à la Technologie et à l’Innovation (CATI) de Maghnia, certifions que :</p>
            <p style="text-align: center; font-weight: bold; font-size: 16px;">{p_genre} {p_nom}</p>
            <p style="text-align: center; font-style: italic;">Qualité : {p_qualite}</p>
            <p>A suivi avec assiduité la formation spécialisée intitulée :</p>
            <p style="text-align: center; font-weight: bold;">« {data['titre_formation']} »</p>
            <p><b>Détails de la session :</b><br>
            Date : {data['date_formation']}<br>
            Durée totale : {data['duree']}<br>
            Lieu : {data['lieu']}<br>
            {formateur_html}</p>
            <p><b>Objectifs et compétences acquises :</b><br>
            {data['objectifs'].replace(chr(10), '<br>')}</p>
            {'<p>L\'évaluation finale a permis de valider les acquis théoriques.</p>' if data['evaluation'] else ''}
            <p>En foi de quoi, la présente attestation lui est délivrée pour servir et valoir ce que de droit.</p>
            <p style="text-align: right;"><b>Le {data['poste_responsable']}</b><br>{data['nom_responsable']}</p>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("---")
        st.header("Téléchargement")
        
        col_dl1, col_dl2 = st.columns(2)
        
        # Sécurisation du nom du fichier pour éviter les espaces et les "nan"
        safe_nom = p_nom.replace(' ', '_').replace('nan', 'Inconnu')

        with col_dl1:
            if st.button("📄 Générer pour ce participant (Word)"):
                docx_buffer = generate_attestation_docx(data, participant)
                nom_fichier = f"Attestation_{safe_nom}.docx"
                st.download_button(
                    label="⬇️ Télécharger Word",
                    data=docx_buffer,
                    file_name=nom_fichier,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                save_to_db(data["reference"], p_nom, data["titre_formation"], data["date_generation"])

        with col_dl2:
            if len(df_participants) > 1:
                if st.button("📦 Générer pour TOUS les participants (ZIP Word)"):
                    zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
                        for idx, row in df_participants.iterrows():
                            docx_buf = generate_attestation_docx(data, row)
                            row_nom = str(row["Nom et Prénom"]) if pd.notna(row["Nom et Prénom"]) else "Inconnu"
                            safe_row_nom = row_nom.replace(' ', '_')
                            nom_fichier = f"Attestation_{safe_row_nom}.docx"
                            zipf.writestr(nom_fichier, docx_buf.getvalue())
                            save_to_db(data["reference"], row_nom, data["titre_formation"], data["date_generation"])
                    
                    zip_buffer.seek(0)
                    st.download_button(
                        label="⬇️ Télécharger le fichier ZIP",
                        data=zip_buffer,
                        file_name="Attestations_CATI.zip",
                        mime="application/zip"
                    )
            else:
                st.info("Ajoutez plusieurs participants pour activer la génération en masse.")
                
        st.markdown("### Conversion PDF")
        st.warning("⚠️ La conversion PDF nécessite Microsoft Word installé sur la machine.")
        if st.button("🔄 Convertir en PDF (Nécessite MS Word)"):
            try:
                temp_docx = "temp.docx"
                temp_pdf = "temp.pdf"
                docx_buf = generate_attestation_docx(data, participant)
                with open(temp_docx, "wb") as f:
                    f.write(docx_buf.getbuffer())
                
                convert(temp_docx, temp_pdf)
                
                with open(temp_pdf, "rb") as f:
                    pdf_bytes = f.read()
                
                st.download_button(
                    label="⬇️ Télécharger PDF",
                    data=pdf_bytes,
                    file_name=f"Attestation_{safe_nom}.pdf",
                    mime="application/pdf"
                )
                
                os.remove(temp_docx)
                os.remove(temp_pdf)
                
            except Exception as e:
                st.error(f"Erreur de conversion PDF : {e}. Assurez-vous que MS Word est installé.")
    else:
        st.warning("Veuillez d'abord remplir et valider le formulaire dans l'onglet 'Saisie des informations'.")

with tab3:
    st.header("Historique des attestations générées")
    history_df = get_history()
    if not history_df.empty:
        st.dataframe(history_df, use_container_width=True)
        csv = history_df.to_csv(index=False).encode('utf-8')
        st.download_button(
            label="⬇️ Exporter l'historique en CSV",
            data=csv,
            file_name='historique_attestations.csv',
            mime='text/csv',
        )
    else:
        st.info("Aucune attestation n'a encore été générée.")
